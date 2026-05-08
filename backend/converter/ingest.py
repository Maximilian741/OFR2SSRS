"""
Multi-file / folder ingest for the Oracle -> SSRS converter.

Public API:
    classify_files(files) -> dict
        Classify a list of (filename, bytes) into useful buckets:
        primary_xml, rdf_binary, sql_files, docs, screenshots, unknown.

    convert_bundle(files) -> dict
        Run the full conversion pipeline given a heterogeneous set of artifacts.
        Returns the same shape as converter.convert(...) plus an "ingest_report"
        key describing what was found / classified.

The goal: let users drop a whole folder of stuff (the Oracle XML, raw .sql
queries, .docx walkthroughs with embedded screenshots, .rdf binaries, loose
.png/.jpg images) and still get *something* useful out -- ideally a real RDL
when an Oracle XML or SQL is present.
"""
from __future__ import annotations

import io
import os
import re
import zipfile
from typing import Any, Dict, List, Optional, Tuple

from .models import DataItem, DataQuery, ParsedReport
from .translators.plsql_to_tsql import translate_report
from .generators.rdl import generate_rdl


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}
_DOC_EXTS = {".txt", ".md", ".rst"}
_SQL_STARTERS = re.compile(
    r"^\s*(SELECT|WITH|FROM|WHERE|JOIN|GROUP\s+BY|ORDER\s+BY|HAVING|UNION|UPDATE|INSERT|DELETE|CREATE)\b",
    re.IGNORECASE,
)
_QUERY_HEADING = re.compile(r"\b(Q_[A-Z0-9_]{1,40})\b")


def _ext(name: str) -> str:
    return os.path.splitext(name or "")[1].lower()


def _decode_text(b: bytes) -> str:
    """Best-effort decode bytes to text."""
    if b is None:
        return ""
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return b.decode(enc)
        except UnicodeDecodeError:
            continue
    return b.decode("utf-8", errors="replace")


def _kind_from_filename(name: str) -> str:
    n = (name or "").lower()
    if "frontend" in n or "front-end" in n or "front_end" in n or "ui" in n:
        return "frontend"
    if "backend" in n or "back-end" in n or "back_end" in n or "server" in n or "db" in n:
        return "backend"
    return "unknown"


# ---------------------------------------------------------------------------
# DOCX extraction (paragraphs + embedded images)
# ---------------------------------------------------------------------------

def _docx_extract(blob: bytes) -> Tuple[str, List[Tuple[str, bytes]]]:
    """Pull paragraph text and embedded media out of a .docx zip.

    Returns (text, [(image_filename, image_bytes), ...]).
    Falls back gracefully if python-docx is missing.
    """
    text_parts: List[str] = []
    images: List[Tuple[str, bytes]] = []

    # Try python-docx first (better paragraph fidelity).
    try:
        import docx  # type: ignore
        d = docx.Document(io.BytesIO(blob))
        for p in d.paragraphs:
            if p.text:
                text_parts.append(p.text)
        # Tables too -- screenshots/queries sometimes live in cells
        for tbl in d.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text:
                            text_parts.append(p.text)
        # Images via the part relationships
        try:
            rels = d.part.rels
            for r_id, rel in rels.items():
                rt = (rel.reltype or "").lower()
                if "image" in rt:
                    try:
                        img_blob = rel.target_part.blob
                        target = getattr(rel, "target_ref", None) or getattr(
                            rel.target_part, "partname", "image"
                        )
                        target_name = os.path.basename(str(target)) or f"image_{r_id}"
                        images.append((target_name, img_blob))
                    except Exception:
                        continue
        except Exception:
            pass
    except Exception:
        # Fallback: parse the docx as a plain zip.
        try:
            with zipfile.ZipFile(io.BytesIO(blob)) as zf:
                for info in zf.infolist():
                    nm = info.filename
                    low = nm.lower()
                    if low == "word/document.xml":
                        try:
                            raw = zf.read(info).decode("utf-8", errors="replace")
                            # crude but effective: strip tags, keep paragraphs by <w:p> spacing
                            raw = re.sub(r"</w:p\s*>", "\n", raw)
                            raw = re.sub(r"<[^>]+>", "", raw)
                            text_parts.append(raw)
                        except Exception:
                            pass
                    elif low.startswith("word/media/") and _ext(low) in _IMAGE_EXTS:
                        try:
                            images.append((os.path.basename(nm), zf.read(info)))
                        except Exception:
                            continue
        except Exception:
            return "", []

    return "\n".join(t for t in text_parts if t), images


def _split_sql_from_text(text: str) -> List[Tuple[str, str]]:
    """Break a free-text doc into a list of (query_name, sql_text) pairs.

    Heuristic: look for headings shaped like Q_PERMIT, Q_ORG, etc. and group
    everything below them. If no such heading exists but the text clearly
    contains SQL statements, return a single anonymous query.
    """
    if not text:
        return []

    lines = text.splitlines()
    sections: List[Tuple[str, List[str]]] = []
    current_name: Optional[str] = None
    current_buf: List[str] = []

    def _flush():
        if current_name and current_buf:
            sections.append((current_name, current_buf[:]))

    for ln in lines:
        m = _QUERY_HEADING.search(ln)
        # A heading line is one where the Q_ marker is essentially the line itself
        if m and len(ln.strip()) <= len(m.group(1)) + 30 and ln.strip().startswith(m.group(1)):
            _flush()
            current_name = m.group(1).upper()
            current_buf = []
            continue
        if current_name is not None:
            current_buf.append(ln)
        else:
            # No active section yet -- buffer SQL-ish text into a default bucket
            current_buf.append(ln)
            if current_name is None and _SQL_STARTERS.match(ln):
                current_name = "Q_DOC"
    _flush()

    out: List[Tuple[str, str]] = []
    for nm, buf in sections:
        joined = "\n".join(buf).strip()
        if not joined:
            continue
        # only keep sections that actually look SQL-ish
        if re.search(r"\bSELECT\b", joined, re.IGNORECASE):
            out.append((nm, joined))

    if not out:
        # Fallback: if the document as a whole has SELECT in it, ship one chunk
        if re.search(r"\bSELECT\b", text, re.IGNORECASE):
            out.append(("Q_DOC", text.strip()))
    return out


# ---------------------------------------------------------------------------
# Classification
# ---------------------------------------------------------------------------

def classify_files(files: List[Tuple[str, bytes]]) -> Dict[str, Any]:
    """Classify a heterogeneous bundle of artifacts.

    Args:
        files: list of (filename, raw_bytes).

    Returns the classification dict described in the module docstring.
    """
    primary_xml: Optional[Tuple[str, bytes]] = None
    rdf_binary: Optional[Tuple[str, bytes]] = None
    sql_files: List[Tuple[str, str]] = []
    pdfs: List[Tuple[str, bytes]] = []
    docs: List[Tuple[str, str]] = []
    screenshots: List[Tuple[str, bytes, str]] = []
    unknown: List[Tuple[str, str]] = []
    summary: List[Dict[str, Any]] = []

    def _note(file: str, category: str, confidence: float, note: str = "") -> None:
        summary.append({
            "file": file,
            "category": category,
            "confidence": round(confidence, 2),
            "note": note,
        })

    for filename, blob in files or []:
        if not filename:
            continue
        ext = _ext(filename)
        base = os.path.basename(filename)

        # --- XML --------------------------------------------------------
        if ext == ".xml":
            is_oracle = False
            note = ""
            try:
                from lxml import etree  # noqa: WPS433
                try:
                    root = etree.fromstring(blob)
                    tag = etree.QName(root.tag).localname.lower() if root.tag else ""
                    dtd = root.get("DTDVersion") or root.get("dtdversion") or ""
                    if tag == "report" and dtd:
                        is_oracle = True
                        note = f"Oracle Reports XML (DTD {dtd})"
                    elif tag == "report":
                        is_oracle = True
                        note = "Oracle Reports XML"
                except Exception as e:
                    note = f"Could not parse XML: {e}"
            except ImportError:
                # lxml missing -- fall back to a string sniff
                head = _decode_text(blob[:2048]).lower()
                if "<report" in head and "dtdversion" in head:
                    is_oracle = True
                    note = "Oracle Reports XML (sniffed)"
                else:
                    note = "lxml missing; could not classify XML"

            if is_oracle and primary_xml is None:
                primary_xml = (filename, blob)
                _note(base, "primary_xml", 0.99, note)
            elif is_oracle:
                # second Oracle XML in the bundle -- keep first, log others
                unknown.append((filename, "duplicate Oracle XML, ignored"))
                _note(base, "unknown", 0.5, "Second Oracle XML in bundle")
            else:
                unknown.append((filename, note or "unknown XML root"))
                _note(base, "unknown", 0.4, note or "XML but not Oracle Reports")
            continue

        # --- PDF (rendered output, used for cross-validation) ---
        if ext == ".pdf":
            pdfs.append((filename, blob))
            _note(base, "pdf", 0.99, "rendered PDF — cross-validation source")
            continue

        # --- RDF binary --------------------------------------------------
        if ext == ".rdf":
            if rdf_binary is None:
                rdf_binary = (filename, blob)
            _note(base, "rdf", 0.95, "Oracle Reports binary (cannot parse natively)")
            continue

        # --- Raw SQL -----------------------------------------------------
        if ext == ".sql":
            sql_text = _decode_text(blob)
            sql_files.append((filename, sql_text))
            _note(base, "sql", 0.95, f"{sum(1 for _ in sql_text.splitlines())} lines")
            continue

        # --- DOCX --------------------------------------------------------
        if ext == ".docx":
            text, embedded = _docx_extract(blob)
            low = base.lower()
            if "sql" in low or "quer" in low:
                # Treat as a SQL doc -- split into (name, sql) tuples.
                pieces = _split_sql_from_text(text)
                if pieces:
                    for nm, sql in pieces:
                        sql_files.append((f"{base}::{nm}", sql))
                    _note(base, "sql", 0.85,
                          f"{len(pieces)} query block(s) extracted from DOCX")
                else:
                    docs.append((filename, text))
                    _note(base, "docs", 0.5,
                          "DOCX named like SQL but no SELECT found; kept as docs")
                # Also harvest any images that happened to be in there
                if embedded:
                    for img_name, img_blob in embedded:
                        screenshots.append((f"{base}::{img_name}", img_blob, "unknown"))
                continue

            kind = _kind_from_filename(base)
            if kind == "frontend":
                if not embedded:
                    docs.append((filename, text))
                    _note(base, "docs", 0.6, "Frontend doc -- no images extracted")
                else:
                    for img_name, img_blob in embedded:
                        screenshots.append((f"{base}::{img_name}", img_blob, "frontend"))
                    _note(base, "screenshot", 0.9,
                          f"{len(embedded)} frontend screenshot(s) extracted")
                continue
            if kind == "backend":
                if not embedded:
                    docs.append((filename, text))
                    _note(base, "docs", 0.6, "Backend doc -- no images extracted")
                else:
                    for img_name, img_blob in embedded:
                        screenshots.append((f"{base}::{img_name}", img_blob, "backend"))
                    _note(base, "screenshot", 0.9,
                          f"{len(embedded)} backend screenshot(s) extracted")
                continue

            # Generic docx
            if text:
                docs.append((filename, text))
            for img_name, img_blob in embedded:
                screenshots.append((f"{base}::{img_name}", img_blob, "unknown"))
            note_bits = []
            if text:
                note_bits.append("text body")
            if embedded:
                note_bits.append(f"{len(embedded)} image(s)")
            _note(base, "docs", 0.7, ", ".join(note_bits) or "empty docx")
            continue

        # --- Standalone images ------------------------------------------
        if ext in _IMAGE_EXTS:
            screenshots.append((filename, blob, _kind_from_filename(base)))
            _note(base, "screenshot", 0.9, f"{len(blob)} bytes")
            continue

        # --- Text/markdown -----------------------------------------------
        if ext in _DOC_EXTS:
            txt = _decode_text(blob)
            docs.append((filename, txt))
            _note(base, "docs", 0.85, f"{len(txt.splitlines())} lines")
            continue

        # --- Fallback ----------------------------------------------------
        unknown.append((filename, "unsupported extension"))
        _note(base, "unknown", 0.3, "unsupported extension")

    return {
        "primary_xml": primary_xml,
        "rdf_binary": rdf_binary,
        "sql_files": sql_files,
        "pdfs": pdfs,
        "docs": docs,
        "screenshots": screenshots,
        "unknown": unknown,
        "category_summary": summary,
    }


# ---------------------------------------------------------------------------
# Bundle conversion
# ---------------------------------------------------------------------------

def _name_from_sql_filename(filename: str) -> str:
    """Derive a query name from a filename like 'Q_PERMIT.sql' or 'docx::Q_ORG'."""
    base = os.path.basename(filename or "")
    # If the synthetic '::' separator is present, prefer the part after it.
    if "::" in base:
        base = base.rsplit("::", 1)[-1]
    name, _ = os.path.splitext(base)
    name = re.sub(r"[^A-Za-z0-9_]", "_", name).strip("_")
    if not name:
        name = "Q_DOC"
    if not name.upper().startswith("Q_"):
        name = "Q_" + name
    return name.upper()


def _build_synthetic_report(sql_files: List[Tuple[str, str]],
                            bundle_label: str = "BUNDLE") -> ParsedReport:
    """Construct a ParsedReport from raw SQL when no Oracle XML is present."""
    report = ParsedReport(name=bundle_label or "BUNDLE", dtd_version="(synthetic)")
    seen: set = set()
    for fname, sql in sql_files:
        if not (sql or "").strip():
            continue
        qname = _name_from_sql_filename(fname)
        candidate = qname
        n = 2
        while candidate in seen:
            candidate = f"{qname}_{n}"
            n += 1
        seen.add(candidate)
        q = DataQuery(name=candidate, sql=sql.strip())
        report.queries.append(q)
    if not report.queries:
        report.warnings.append("No SQL queries could be extracted from the bundle.")
    else:
        report.warnings.append(
            "Synthetic report: built from raw SQL artifacts (no Oracle XML provided)."
        )
    return report


def _ingest_report_dict(classification: Dict[str, Any]) -> Dict[str, Any]:
    """Turn the raw classification into a JSON-friendly summary for the UI."""
    primary_xml = classification.get("primary_xml")
    rdf_binary = classification.get("rdf_binary")
    return {
        "primary_xml": primary_xml[0] if primary_xml else None,
        "rdf_binary": rdf_binary[0] if rdf_binary else None,
        "sql_files": [
            {"file": fn, "lines": len((txt or "").splitlines()),
             "size": len(txt or "")}
            for fn, txt in classification.get("sql_files", [])
        ],
        "docs": [
            {"file": fn, "chars": len(txt or "")}
            for fn, txt in classification.get("docs", [])
        ],
        "screenshots": [
            {"file": fn, "kind": kind, "size": len(blob or b"")}
            for fn, blob, kind in classification.get("screenshots", [])
        ],
        "unknown": [
            {"file": fn, "reason": reason}
            for fn, reason in classification.get("unknown", [])
        ],
        "category_summary": list(classification.get("category_summary", [])),
        "totals": {
            "xml": 1 if primary_xml else 0,
            "rdf": 1 if rdf_binary else 0,
            "sql": len(classification.get("sql_files", [])),
            "docs": len(classification.get("docs", [])),
            "screenshots": len(classification.get("screenshots", [])),
            "unknown": len(classification.get("unknown", [])),
        },
    }


def convert_bundle(files: List[Tuple[str, bytes]]) -> Dict[str, Any]:
    """Top-level entry: classify everything, then convert what we can.

    Returns the same shape as converter.convert(...) plus an "ingest_report"
    key. If no convertible artifacts exist, returns
    {"error": "no_convertible_artifacts", "ingest_report": {...}}.
    """
    classification = classify_files(files or [])
    ingest = _ingest_report_dict(classification)

    primary_xml = classification.get("primary_xml")
    sql_files = classification.get("sql_files") or []

    # Path 1: real Oracle XML present -- run the regular pipeline.
    if primary_xml:
        from . import convert as _convert  # local import to avoid cycles
        from .cross_validate import cross_validate
        from .parsers.oracle_xml import parse_oracle_xml
        try:
            data = _convert(primary_xml[1])
        except Exception as e:
            return {
                "error": f"convert_failed: {e}",
                "ingest_report": ingest,
            }
        data["ingest_report"] = ingest
        # Cross-validation against any supporting artifacts in the bundle
        try:
            parsed_for_xv = parse_oracle_xml(primary_xml[1])
            data["cross_validation"] = cross_validate(
                parsed_for_xv, classification, data.get("bursting"))
        except Exception as _e:
            data["cross_validation"] = {"error": f"{type(_e).__name__}: {_e}"}
        return data

    # Path 2: no XML, but we have SQL -- build a synthetic report.
    if sql_files:
        label = "BUNDLE"
        # Try to guess a nicer label from filenames
        for fn, _ in sql_files:
            base = os.path.basename(fn).split("::", 1)[0]
            stem, _ = os.path.splitext(base)
            if stem:
                label = re.sub(r"[^A-Za-z0-9_]", "_", stem).upper()
                break
        report = _build_synthetic_report(sql_files, bundle_label=label)
        try:
            translate_report(report)
        except Exception as e:
            report.warnings.append(f"Translator error: {e}")
        try:
            rdl_xml = generate_rdl(report)
        except Exception as e:
            rdl_xml = ""
            report.warnings.append(f"RDL generation error: {e}")
        return {
            "report": report.to_dict(),
            "rdl_xml": rdl_xml,
            "oracle_xml": "",
            "mockup_html": "<em>No Oracle layout available -- synthetic report from SQL only.</em>",
            "ingest_report": ingest,
        }

    # Path 3: nothing convertible.
    return {
        "error": "no_convertible_artifacts",
        "ingest_report": ingest,
    }


__all__ = ["classify_files", "convert_bundle"]

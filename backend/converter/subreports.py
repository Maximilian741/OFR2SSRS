"""
Sub-report (drill-through) detection + composition.

Oracle Reports lets a field be a hyperlink that opens ANOTHER report,
passing parameter values in the URL query string:

    /reports/rwservlet?SERVER&destype=cache&desformat=PDF
        &report=CHILD_REPORT.rep&P_FOO=bar

The parent XML declares this via:

  * a <userParameter> whose name commonly starts with P_AS_PATH or
    P_ENVELOPE / P_URL_* and whose initialValue NAMES the child report
  * one or more PL/SQL formulas (CF_URL_*, CP_URL_*) that concatenate
    the parameters into the URL string
  * a layout field carrying <webSettings hyperlink="&CF_URL_X">

This module:

  1. Scans a ParsedReport for those signatures and returns a
     normalized list of { child_name, link_text, parent_field,
     url_template, bind_params } so the frontend can surface them.
  2. Synthesizes a minimal RDL stub for a child report when only
     artifacts (SQL text in a .docx/.sql, optional screenshots) are
     available -- so the user can drag a child .rdl onto the SSRS
     server even without the child's Oracle XML.

Everything is purely pattern-based: NOTHING is hard-coded for any
specific report, parameter, or organization.
"""
from __future__ import annotations
import os
import re
import zipfile
from typing import Any, Dict, Iterable, List, Optional, Tuple


# ---------------------------------------------------------------------------
# Detection
# ---------------------------------------------------------------------------

# Parameter names that commonly NAME a child report or carry the child
# report's URL/path. Drill-through reports use one of these.
_CHILD_REPORT_PARAM_HINTS = (
    "P_AS_PATH", "P_ENVELOPE", "P_DRILLDOWN", "P_DRILL", "P_SUBREPORT",
    "P_CHILD_REPORT", "P_NESTED",
)

# Parameter names that carry a base URL to a child SSRS/Oracle reports
# endpoint -- generic, just a hint that drill-through composition is
# happening somewhere.
_URL_PARAM_HINTS = (
    "P_URL_", "P_REPORT_SERVER", "P_URL", "P_REP_URL",
)

# Regex that finds "report=<NAME>" tokens inside a URL-builder formula.
_REPORT_NAME_RE = re.compile(
    r"report\s*=\s*[\"']?([A-Za-z][A-Za-z0-9_]+?)(?:\.rep)?\b",
    re.IGNORECASE,
)

# Hyperlink markers in the Oracle XML layout.
_HYPERLINK_RE = re.compile(
    r'hyperlink\s*=\s*"&amp;?([A-Za-z_][A-Za-z0-9_]*)"',
    re.IGNORECASE,
)


def _norm(s: str) -> str:
    return (s or "").upper().strip()


def _all_param_names(report) -> List[str]:
    return [p.name for p in getattr(report, "parameters", []) or [] if p.name]


def _param_lookup(report, name: str):
    n = _norm(name)
    for p in getattr(report, "parameters", []) or []:
        if _norm(p.name) == n:
            return p
    return None


def _formulas(report) -> List[Any]:
    return list(getattr(report, "formulas", []) or [])


def _raw_xml(report) -> str:
    """Best-effort recovery of the raw source XML for the layout-level
    scans (hyperlink= attributes live in the XML but may not survive
    parsing into our model)."""
    return getattr(report, "raw_xml", "") or ""


def _bind_names(body: str) -> List[str]:
    return [m.group(1) for m in re.finditer(r":([A-Za-z_][A-Za-z0-9_]*)", body or "")]


def _amp_refs(body: str) -> List[str]:
    return [m.group(1) for m in re.finditer(r"&([A-Za-z_][A-Za-z0-9_]*)", body or "")]


def detect_subreport_links(report) -> List[Dict[str, Any]]:
    """Return a normalized list of detected drill-through links.

    Each item:
        {
          "child_name":    "CHILD_REPORT",     # or None when unknown
          "link_text":     "[Permittee]",      # the textual hint
          "parent_field":  "F_CP_CHILD",       # the layout field that
                                               # carries webSettings hyperlink
          "url_formula":   "CF_URL_ENVELOPE",  # formula building the URL
          "bind_params":   ["P_PERM_NUM", ...],# parent params forwarded
        }

    Multiple hyperlink markers in the same report produce multiple
    entries. An empty list means no drill-through was detected.
    """
    links: List[Dict[str, Any]] = []
    raw = _raw_xml(report)
    params = _all_param_names(report)
    param_upper = {p.upper() for p in params}

    # 1. Pull every formula source -> body
    formula_bodies: Dict[str, str] = {}
    for f in _formulas(report):
        fname = (getattr(f, "name", "") or "").strip()
        body = getattr(f, "plsql_body", "") or getattr(f, "body", "") or ""
        if fname:
            formula_bodies[fname] = body

    # 2. Pull the candidate child-report name from any *_PATH /
    #    *_ENVELOPE / *_REPORT parameter that has an initialValue.
    candidate_child_names: List[str] = []
    for p in getattr(report, "parameters", []) or []:
        name = _norm(p.name)
        if any(h in name for h in _CHILD_REPORT_PARAM_HINTS):
            iv = (getattr(p, "initial_value", "") or "").strip()
            if iv and re.match(r"^[A-Za-z][A-Za-z0-9_]+$", iv):
                candidate_child_names.append(iv)

    # Also: scan formula bodies for any "report=<NAME>.rep" substring.
    formula_report_refs: Dict[str, str] = {}
    for fname, body in formula_bodies.items():
        m = _REPORT_NAME_RE.search(body)
        if m:
            formula_report_refs[fname] = m.group(1)

    # 3. Scan the raw XML for every hyperlink= attribute. Each one
    #    points at a placeholder/formula whose body has the URL.
    hyperlink_hits = list(_HYPERLINK_RE.finditer(raw))
    seen_url_sources: set = set()

    for m in hyperlink_hits:
        url_source = m.group(1)  # e.g. CP_URL_ALL_ENVELOPE or CF_URL_ENVELOPE
        if url_source in seen_url_sources:
            continue
        seen_url_sources.add(url_source)

        # Find which formula populates this placeholder/source.
        body = formula_bodies.get(url_source, "")
        # If the source is a CP_ placeholder, look for a CF_ formula
        # that assigns to it (Oracle uses ":CP_NAME := ..." pattern).
        if not body:
            for fname, fbody in formula_bodies.items():
                if re.search(
                    rf":{re.escape(url_source)}\s*:?=", fbody, re.IGNORECASE,
                ):
                    body = fbody
                    break

        # Extract child report name from the URL body (if any).
        child_name = None
        rm = _REPORT_NAME_RE.search(body)
        if rm:
            child_name = rm.group(1)
        # Fall back to candidate_child_names (P_AS_PATH initialValue).
        if not child_name and candidate_child_names:
            child_name = candidate_child_names[0]

        # Forwarded params: any bind variable in the URL body that's a
        # report parameter (filters out destype, desformat, etc.).
        binds = _bind_names(body)
        forwarded = [
            b for b in binds
            if b.upper() in param_upper
            and not b.upper().startswith(("P_DEST", "P_DES"))
        ]
        # De-duplicate while preserving order.
        seen = set()
        forwarded = [b for b in forwarded if not (b in seen or seen.add(b))]

        # Try to pull the link's text hint: the static <text> block
        # often labels which field is the link target.
        link_text = None
        # Look for a label-like name patten: "X is a hyperlink to Y"
        tm = re.search(
            rf'\[([A-Za-z][A-Za-z0-9_]*)\]\s+is\s+a\s+hyperlink',
            raw,
            re.IGNORECASE,
        )
        if tm:
            link_text = tm.group(1)

        # Parent field name: the field whose webSettings carries
        # this hyperlink. Inspect a window around the match.
        parent_field = None
        # Look BACKWARD from the hyperlink position for the nearest
        # <field name="X" ...> declaration.
        upto = raw[: m.start()]
        fm = list(re.finditer(r'<field\s+name="([^"]+)"', upto, re.IGNORECASE))
        if fm:
            parent_field = fm[-1].group(1)

        links.append({
            "child_name": child_name,
            "link_text": link_text,
            "parent_field": parent_field,
            "url_formula": url_source,
            "bind_params": forwarded,
        })

    return _dedupe_by_child(links)


def _csv_union(a: Optional[str], b: Optional[str]) -> Optional[str]:
    """Merge two comma-separated string fields, preserving order and
    skipping duplicates. Returns None when both inputs are empty."""
    parts: List[str] = []
    seen: set = set()
    for src in (a, b):
        if not src:
            continue
        for piece in str(src).split(","):
            p = piece.strip()
            if p and p not in seen:
                seen.add(p)
                parts.append(p)
    return ", ".join(parts) if parts else None


def _dedupe_by_child(links: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Collapse multiple link entries that resolve to the SAME child
    report into one. Oracle parents commonly hyperlink to one child
    from several layout fields via different URL-builder formulas; the
    artifact store and the build endpoint already key on child_name,
    so showing a separate UI card per layout field is just confusing.

    When merging, parent_field / url_formula are joined as a comma-
    separated list (preserving order, de-duped), bind_params are
    unioned, and the first non-empty link_text wins. Entries whose
    child_name could not be resolved are kept separate -- we can't
    prove they point at the same child.

    Generic: groups purely by the resolved child name; no per-report
    knowledge.
    """
    out: List[Dict[str, Any]] = []
    by_child: Dict[str, Dict[str, Any]] = {}
    for ln in links:
        cn = ln.get("child_name")
        if not cn:
            out.append(ln)
            continue
        if cn not in by_child:
            merged = dict(ln)
            by_child[cn] = merged
            out.append(merged)
            continue
        cur = by_child[cn]
        cur["parent_field"] = _csv_union(cur.get("parent_field"),
                                         ln.get("parent_field"))
        cur["url_formula"] = _csv_union(cur.get("url_formula"),
                                        ln.get("url_formula"))
        if not cur.get("link_text"):
            cur["link_text"] = ln.get("link_text")
        seen: set = set()
        merged_binds: List[str] = []
        for b in (list(cur.get("bind_params") or [])
                  + list(ln.get("bind_params") or [])):
            if b and b not in seen:
                seen.add(b)
                merged_binds.append(b)
        cur["bind_params"] = merged_binds
    return out


def is_drillthrough_only(report) -> bool:
    """True when the report has hyperlink-style child-report links but
    NO real bursting markers (per-row distribution path/email).

    Used to suppress the bursting flag for drill-through-only reports
    like a permit/letter report, which were previously mis-classified as bursting
    because they declare P_AS_PATH (a URL builder, not a destination).
    """
    if not detect_subreport_links(report):
        return False
    # Heuristic: real bursting also has an email/contact field or
    # multiple per-row output paths. If we don't see those, it's just
    # a drill-through link.
    raw = _raw_xml(report)
    has_email = bool(re.search(
        r"(EMAIL|MAIL|RECIPIENT)_(ADDR|ADDRESS|TO)",
        raw, re.IGNORECASE,
    ))
    has_distribution = bool(re.search(
        r"<distribution\b|<destination\b", raw, re.IGNORECASE,
    ))
    return not (has_email or has_distribution)


# ---------------------------------------------------------------------------
# Composition -- synthesize a minimal RDL for a child report whose XML
# we DON'T have, working from whatever artifacts the user uploaded
# (SQL text in .sql or .docx, optional screenshots, etc.)
# ---------------------------------------------------------------------------

def _extract_text_from_docx(path: str) -> str:
    """Pull paragraph text out of a .docx.

    Uses python-docx so the runs inside a paragraph are concatenated
    correctly. Word splits edited text into many <w:r> runs; a naive
    tag-strip drops the run boundaries AND turns every paragraph break
    into a newline -- which slices tokens like ``:P_PARAM`` apart at the
    colon and corrupts the extracted SQL (the child sub-report then gets
    an unrunnable CommandText). python-docx joins runs per paragraph, so
    bind variables and identifiers stay intact.
    """
    try:
        import docx  # python-docx -- already a dependency (see ingest.py)
        document = docx.Document(path)
    except Exception:
        # Fallback: crude tag-strip. Better than nothing if python-docx
        # cannot open the file, though run boundaries may be lost.
        try:
            with zipfile.ZipFile(path) as z:
                data = z.read("word/document.xml").decode("utf-8", "replace")
        except Exception:
            return ""
        text = re.sub(r"</w:p>", "\n", data)
        text = re.sub(r"<[^>]+>", "", text)
        return re.sub(r"\n{2,}", "\n", text).strip()

    lines = [p.text for p in document.paragraphs]
    # SQL is sometimes pasted into a table cell -- pull that text too.
    for table in getattr(document, "tables", []) or []:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    lines.append(p.text)
    return "\n".join(lines).strip()


def _sql_from_artifact(path: str) -> str:
    """Best-effort SQL extraction from an artifact file."""
    name = path.lower()
    if name.endswith(".sql") or name.endswith(".txt"):
        try:
            return open(path, "r", encoding="utf-8", errors="replace").read()
        except Exception:
            return ""
    if name.endswith(".docx"):
        return _extract_text_from_docx(path)
    return ""


def _select_columns(sql: str) -> List[str]:
    """Return the column-name list for the first SELECT in the SQL.

    Pure regex; handles aliases (AS NAME or trailing NAME) and bare
    table.col references. Anything we can't parse cleanly is skipped.
    """
    if not sql:
        return []
    m = re.search(r"\bSELECT\b(.+?)\bFROM\b", sql, re.IGNORECASE | re.DOTALL)
    if not m:
        return []
    body = m.group(1)
    # Split on top-level commas.
    parts: List[str] = []
    cur: List[str] = []
    depth = 0
    for ch in body:
        if ch == "(":
            depth += 1
        elif ch == ")":
            depth -= 1
        if ch == "," and depth == 0:
            parts.append("".join(cur))
            cur = []
        else:
            cur.append(ch)
    if cur:
        parts.append("".join(cur))

    cols: List[str] = []
    for raw in parts:
        item = raw.strip().rstrip(",").strip()
        if not item:
            continue
        # explicit alias: "... AS NAME"
        am = re.search(
            r"\bAS\s+([A-Za-z_][A-Za-z0-9_]*)\s*$", item, re.IGNORECASE,
        )
        if am:
            cols.append(am.group(1))
            continue
        # bare TABLE.COL  ->  COL
        bm = re.match(
            r"^\s*[A-Za-z_][A-Za-z0-9_]*\.([A-Za-z_][A-Za-z0-9_]*)\s*$", item,
        )
        if bm:
            cols.append(bm.group(1))
            continue
        # bare COL
        single = re.match(r"^\s*([A-Za-z_][A-Za-z0-9_]*)\s*$", item)
        if single:
            cols.append(single.group(1))
            continue
        # trailing identifier alias: "func(...) NAME"
        tm = re.search(
            r"(?:\)|[A-Za-z0-9_])\s+([A-Za-z_][A-Za-z0-9_]*)\s*$", item,
        )
        if tm:
            cols.append(tm.group(1))
    # Dedupe (preserve order)
    seen = set()
    return [c for c in cols if not (c in seen or seen.add(c))]


def _bind_params_in_sql(sql: str) -> List[str]:
    """Return bind variable names referenced in the SQL, in order."""
    seen = set()
    out = []
    for m in re.finditer(r":([A-Za-z_][A-Za-z0-9_]*)", sql or ""):
        n = m.group(1)
        if n not in seen:
            seen.add(n)
            out.append(n)
    return out


def _lexical_refs_in_sql(sql: str) -> List[str]:
    """Return Oracle Reports lexical-reference names (``&P_CRITERIA`` style) in
    the SQL, in order. A lexical splices a raw SQL FRAGMENT (a whole WHERE /
    ORDER BY clause) at parse time -- SSRS has no direct equivalent. We surface
    them so the build can tell the user exactly which knob controls filtering.
    Skips the XML entity escapes (&amp;/&lt;/&gt;/&quot;/&apos;)."""
    seen, out = set(), []
    for m in re.finditer(r"&([A-Za-z_][A-Za-z0-9_]*)", sql or ""):
        n = m.group(1)
        if n.lower() in ("amp", "lt", "gt", "quot", "apos", "nbsp"):
            continue
        if n.upper() not in seen:
            seen.add(n.upper())
            out.append(n)
    return out


def _id_columns_in_sql(sql: str) -> List[str]:
    """Return qualified ``ALIAS.Something_Id`` columns referenced in the SQL,
    in order, de-duped. These are the candidate join keys a drill-through
    filter would target (e.g. O.Org_Id, SA.Site_Id, S.Site_Id) -- surfaced as
    a hint so the user can wire the forwarded params into the WHERE clause.
    Purely structural (regex on ``alias.colname`` ending in _id), no report-
    specific names."""
    seen, out = set(), []
    for m in re.finditer(r"\b([A-Za-z_]\w*\.[A-Za-z_]\w*?_id)\b", sql or "", re.IGNORECASE):
        col = m.group(1)
        if col.upper() not in seen:
            seen.add(col.upper())
            out.append(col)
    return out


def forwarded_drillthrough_params(report, child_name: str) -> List[str]:
    """Parameter NAMES the parent report forwards to ``child_name`` through a
    drill-through URL formula (e.g. ['P_ORG_ID', 'P_SITE_ID']).

    These come from the ``'&PARAM=' || :source`` pairs in the parent's
    CF_URL_*/CP_URL_* formula -- the SAME pairs the parent's <Drillthrough>
    emits. The child RDL MUST declare every one of them or SSRS raises
    "the parameter 'P_ORG_ID' is not declared" the instant the link is
    clicked. Generic: parsed from formula text, nothing report-specific.
    """
    if report is None or not child_name:
        return []
    try:
        from .generators.rdl import _extract_url_params
    except Exception:  # noqa: BLE001
        return []
    target = _norm(child_name)
    out, seen = [], set()
    for ln in detect_subreport_links(report):
        if _norm(ln.get("child_name") or "") != target:
            continue
        names = {x.strip().upper() for x in (ln.get("url_formula") or "").split(",")}
        for pname, _src in _extract_url_params(report, names):
            if pname.upper() not in seen:
                seen.add(pname.upper())
                out.append(pname)
    return out


def _trim_to_first_statement(sql: str) -> str:
    """Cut a free-text SQL blob down to just the first SELECT statement.

    Report .docx/.sql artifacts often paste the query AND the report's
    PL/SQL (After_Param_Form, CF_* formulas, packages) into one blob with
    no ';' terminating the SELECT. Stop at the earliest of: a ';', a lone
    SQL*Plus '/' line, or a line that STARTS a PL/SQL block
    (FUNCTION / PROCEDURE / PACKAGE / DECLARE / BEGIN). Generic --
    keyword/structure based, never report-specific.
    """
    if not sql:
        return sql
    m = re.search(r"\bSELECT\b", sql, re.IGNORECASE)
    if not m:
        return sql.strip()
    body = sql[m.start():]
    cut = len(body)
    semi = body.find(";")
    if semi != -1:
        cut = min(cut, semi)
    kw = re.search(r"(?im)^[ \t]*(FUNCTION|PROCEDURE|PACKAGE|DECLARE|BEGIN)\b", body)
    if kw:
        cut = min(cut, kw.start())
    slash = re.search(r"(?m)^[ \t]*/[ \t]*$", body)
    if slash:
        cut = min(cut, slash.start())
    return body[:cut].strip()


def compose_subreport_rdl(child_name: str,
                          artifacts: Iterable[str],
                          parent_param_names: Optional[Iterable[str]] = None,
                          drillthrough_params: Optional[Iterable[str]] = None
                          ) -> Dict[str, Any]:
    """Synthesize a minimal RDL stub for a child report from artifacts.

    Parameters
    ----------
    child_name : str
        The child report's identifier (e.g. "CHILD_REPORT"). Used as
        the <DataSet Name="..."> and the file name.
    artifacts : iterable of file paths
        Any uploaded SQL/DOCX/TXT files. We pull SQL from the first
        artifact that yields a SELECT; columns are inferred from it.
    parent_param_names : optional iterable
        Names of the parent report's parameters. Any of those that
        ALSO appear as bind variables in the child SQL become
        ReportParameters in the child stub (so the parent's
        drill-through can pass them through).

    Returns
    -------
    dict with:
        "rdl_xml":   str         (a complete RDL string, ready to drop)
        "issues":    list[str]   (anything we couldn't infer)
        "fields":    list[str]   (columns we surfaced)
        "binds":     list[str]   (bind variables found in SQL)
        "sql":       str         (the SQL we extracted)
    """
    # 1. Locate SQL among the artifacts.
    sql = ""
    for a in artifacts or []:
        if not os.path.isfile(a):
            continue
        candidate = _sql_from_artifact(a)
        # Quick gate: must contain SELECT ... FROM
        if re.search(r"\bSELECT\b.+?\bFROM\b", candidate or "",
                     re.IGNORECASE | re.DOTALL):
            sql = candidate
            break

    # Trim free-form artifact text down to the first SQL statement so the
    # child <CommandText> is runnable. A sql.docx commonly carries a
    # title line ("<REPORT> sql statements"), a query-block label
    # ("Q_ADDRESS"), and several query blocks around the SELECT -- none
    # of which belong in CommandText.
    if sql:
        sql = _trim_to_first_statement(sql)

    fields = _select_columns(sql)
    binds = _bind_params_in_sql(sql)
    lexicals = _lexical_refs_in_sql(sql)
    # Neutralize Oracle lexical refs (&P_CRITERIA) so the stub CommandText is
    # valid SQL -- the same rule the real generator applies. A raw "&NAME"
    # reaches Oracle as a syntax error; a comment keeps the statement runnable.
    if sql and lexicals:
        sql = re.sub(r"&([A-Za-z_][A-Za-z0-9_]*)",
                     r"/* lexical ref &\1 -- wire as dynamic WHERE at deploy time */",
                     sql)

    # 2. Build the RDL XML. Keep it minimal but schema-valid so the
    #    user can upload it to SSRS without errors.
    parent_param_set = {p.upper() for p in (parent_param_names or [])}
    dt_params = [p for p in (drillthrough_params or []) if p]
    bind_upper = {b.upper() for b in binds}
    # Every parameter the child declares: SQL binds + the parent's forwarded
    # drill-through params (which the child SQL may not bind directly).
    dt_only = [p for p in dt_params if p.upper() not in bind_upper]
    forwarded = [b for b in binds if b.upper() in parent_param_set] + dt_only
    safe = lambda s: re.sub(r"[^A-Za-z0-9_]", "_", s)
    cname = safe(child_name) or "SubReport"
    # When no SQL was extracted from the artifacts, emit a runnable
    # placeholder query rather than a bare comment. A SQL comment alone
    # is not a valid statement on Oracle (ORA-00900) and SSRS will fail
    # the dataset refresh on upload. ``SELECT 'PLACEHOLDER' AS PLACEHOLDER
    # FROM DUAL`` is valid on Oracle; the FROM DUAL clause is harmless on
    # SQL Server-bound stubs because Oracle sub-reports are the common
    # case for this code path. The user can replace the query in Report
    # Builder once they have the real SQL.
    placeholder_sql = (
        f"-- Placeholder query for sub-report {child_name}. Replace with "
        f"the real SELECT once available.\n"
        f"SELECT 'PLACEHOLDER' AS PLACEHOLDER FROM DUAL"
    )
    safe_sql = (sql or placeholder_sql).replace("&", "&amp;") \
                                       .replace("<", "&lt;")

    field_xml = "\n".join(
        f'        <Field Name="{safe(c)}">\n'
        f'          <DataField>{safe(c)}</DataField>\n'
        f'          <rd:TypeName>System.String</rd:TypeName>\n'
        f'        </Field>'
        for c in (fields or ["PLACEHOLDER"])
    )

    qparam_xml = ""
    if binds:
        qparams = "\n".join(
            f'        <QueryParameter Name=":{b}">\n'
            f'          <Value>=Parameters!{safe(b)}.Value</Value>\n'
            f'        </QueryParameter>'
            for b in binds
        )
        qparam_xml = f"      <QueryParameters>\n{qparams}\n      </QueryParameters>\n"

    # ReportParameters: SQL binds + drill-through-only forwarded params.
    # EVERY one gets <Nullable>true</Nullable> + a =Nothing <DefaultValue> so
    # "Refresh Fields" / dataset refresh NEVER pops the "Define Query
    # Parameters" prompt -- the load-bearing invariant (a param with no default
    # is exactly what triggers that prompt). Drill-through-only params are
    # Hidden (the parent sets them; the standalone user shouldn't see them).
    def _rparam(name: str, hidden: bool) -> str:
        return (
            f'    <ReportParameter Name="{safe(name)}">\n'
            f'      <DataType>String</DataType>\n'
            f'      <Nullable>true</Nullable>\n'
            f'      <DefaultValue>\n'
            f'        <Values>\n'
            f'          <Value>=Nothing</Value>\n'
            f'        </Values>\n'
            f'      </DefaultValue>\n'
            f'      <AllowBlank>true</AllowBlank>\n'
            f'      <Prompt>{name}</Prompt>\n'
            + ('      <Hidden>true</Hidden>\n' if hidden else '')
            + f'    </ReportParameter>'
        )
    rparam_xml = ""
    all_params = [(b, False) for b in binds] + [(p, True) for p in dt_only]
    if all_params:
        rparams = "\n".join(_rparam(n, h) for n, h in all_params)
        rparam_xml = f'  <ReportParameters>\n{rparams}\n  </ReportParameters>\n'

    rdl = (
        '<?xml version="1.0" encoding="utf-8"?>\n'
        '<Report xmlns="http://schemas.microsoft.com/sqlserver/reporting/2008/01/reportdefinition" '
        'xmlns:rd="http://schemas.microsoft.com/SQLServer/reporting/reportdesigner">\n'
        '  <DataSources>\n'
        '    <DataSource Name="SharedDataSource">\n'
        '      <DataSourceReference>SharedDataSource</DataSourceReference>\n'
        '      <rd:SecurityType>None</rd:SecurityType>\n'
        '      <rd:DataSourceID>00000000-0000-0000-0000-000000000001</rd:DataSourceID>\n'
        '    </DataSource>\n'
        '  </DataSources>\n'
        '  <DataSets>\n'
        f'    <DataSet Name="{cname}">\n'
        '      <Query>\n'
        '        <DataSourceName>SharedDataSource</DataSourceName>\n'
        f'        <CommandText>{safe_sql}</CommandText>\n'
        f'{qparam_xml}'
        '      </Query>\n'
        '      <Fields>\n'
        f'{field_xml}\n'
        '      </Fields>\n'
        '    </DataSet>\n'
        '  </DataSets>\n'
        f'{rparam_xml}'
        '  <Body>\n'
        '    <ReportItems>\n'
        f'      <Textbox Name="Tb_SubreportTitle">\n'
        '        <Paragraphs>\n'
        '          <Paragraph>\n'
        '            <TextRuns>\n'
        f'              <TextRun><Value>="{child_name} -- Sub-report (composed from artifacts; refine layout in Report Builder)"</Value>\n'
        '                <Style><FontSize>11pt</FontSize><FontWeight>Bold</FontWeight></Style>\n'
        '              </TextRun>\n'
        '            </TextRuns>\n'
        '          </Paragraph>\n'
        '        </Paragraphs>\n'
        '        <Top>0.1in</Top><Left>0.1in</Left>\n'
        '        <Width>7.3in</Width><Height>0.3in</Height>\n'
        '        <CanGrow>true</CanGrow>\n'
        '      </Textbox>\n'
        '    </ReportItems>\n'
        '    <Height>2in</Height>\n'
        '    <Style />\n'
        '  </Body>\n'
        '  <Width>7.5in</Width>\n'
        '  <Page>\n'
        '    <PageHeader><Height>0.25in</Height><PrintOnFirstPage>true</PrintOnFirstPage><PrintOnLastPage>true</PrintOnLastPage></PageHeader>\n'
        '    <PageFooter><Height>0.25in</Height><PrintOnFirstPage>true</PrintOnFirstPage><PrintOnLastPage>true</PrintOnLastPage></PageFooter>\n'
        '    <PageHeight>11in</PageHeight><PageWidth>8.5in</PageWidth>\n'
        '    <LeftMargin>0.5in</LeftMargin><RightMargin>0.5in</RightMargin>\n'
        '    <TopMargin>0.5in</TopMargin><BottomMargin>0.5in</BottomMargin>\n'
        '  </Page>\n'
        '  <Code></Code>\n'
        '  <Language>en-US</Language>\n'
        '</Report>\n'
    )

    issues: List[str] = []
    if not sql:
        issues.append("no SQL was found in any uploaded artifact -- "
                      "child RDL has a placeholder CommandText.")
    if not fields:
        issues.append("could not infer column list from SQL -- "
                      "PLACEHOLDER field emitted; refine in Report Builder.")
    issues.extend(_drillthrough_reconciliation_notes(sql, dt_params))
    if parent_param_set and not forwarded:
        issues.append("parent has parameters but none match SQL binds -- "
                      "drill-through will not forward values.")

    return {
        "rdl_xml": rdl,
        "issues": issues,
        "fields": fields,
        "binds": binds,
        "forwarded_params": forwarded,
        "sql": sql,
    }


# ---------------------------------------------------------------------------
# Rich sub-report build
#
# Routes a child report's artifacts through the SAME pipeline the main report
# uses, so the child gets a full RDL *and* an HTML mockup preview (the
# minimal stub above is only the last-resort fallback). Accepts ANY artifact:
#
#   * an Oracle Reports XML   -> full convert() (real RDL + real mockup)
#   * an existing .rdl        -> used as-is; preview derived from its dataset
#   * SQL in .sql/.docx/.txt  -> synthesized report -> real generator + mockup
#   * nothing parseable       -> stub fallback with a "drop more" note
#
# Generic: every decision derives from artifact CONTENT/STRUCTURE, never from
# a specific report name, parameter, or field list.
# ---------------------------------------------------------------------------

def _read_artifacts(paths: Iterable[str]) -> List[Tuple[str, bytes]]:
    out: List[Tuple[str, bytes]] = []
    for p in paths or []:
        try:
            if os.path.isfile(p):
                with open(p, "rb") as fh:
                    out.append((os.path.basename(p), fh.read()))
        except Exception:
            continue
    return out


def _looks_like_oracle_xml(name: str, blob: bytes) -> bool:
    try:
        head = blob[:4096].decode("utf-8", "replace").lower()
    except Exception:
        return False
    if "<report" not in head:
        return False
    # An Oracle Reports XML carries a DTDVersion and/or the data/layout tags
    # the parser understands. An SSRS .rdl also has <Report> but a very
    # different namespace -- keep them apart so .rdl files go down the
    # pass-through branch instead of the parser branch.
    if "reportdefinition" in head:
        return False
    return ("dtdversion" in head or "<datasource" in head
            or "<userparameter" in head or "<layout" in head)


def _looks_like_rdl(name: str, blob: bytes) -> bool:
    low = (name or "").lower()
    try:
        head = blob[:4096].decode("utf-8", "replace").lower()
    except Exception:
        head = ""
    if low.endswith(".rdl"):
        return True
    return "<report" in head and "reportdefinition" in head


def _unescape_xml(s: str) -> str:
    return (s.replace("&lt;", "<").replace("&gt;", ">")
             .replace("&quot;", '"').replace("&apos;", "'")
             .replace("&amp;", "&"))


def _report_from_rdl(rdl_text: str, child_name: str):
    """Best-effort ParsedReport from an existing RDL, purely so we can render
    a preview mockup. Picks the dataset with the most <Field>s as the main
    block. Namespace-agnostic (regex) so it works for any RDL schema version.
    """
    from .models import ParsedReport, DataQuery, DataItem
    rep = ParsedReport(name=child_name or "SubReport", dtd_version="(from RDL)")
    datasets = re.findall(
        r"<DataSet\b[^>]*\bName=\"([^\"]+)\"(.*?)</DataSet>",
        rdl_text, re.DOTALL | re.IGNORECASE,
    )
    best = None  # (field_count, ds_name, block)
    for ds_name, block in datasets:
        fcount = len(re.findall(r"<Field\b[^>]*\bName=\"[^\"]+\"", block,
                                re.IGNORECASE))
        if best is None or fcount > best[0]:
            best = (fcount, ds_name, block)
    if best:
        _, ds_name, block = best
        cmd_m = re.search(r"<CommandText>(.*?)</CommandText>", block,
                          re.DOTALL | re.IGNORECASE)
        sql = _unescape_xml(cmd_m.group(1)).strip() if cmd_m else ""
        items = [
            DataItem(name=fn, label=fn.replace("_", " ").title())
            for fn in re.findall(r"<Field\b[^>]*\bName=\"([^\"]+)\"", block,
                                 re.IGNORECASE)
        ]
        rep.queries.append(DataQuery(name=ds_name, sql=sql, tsql=sql, items=items))
    return rep


def _inject_report_parameters(rdl_text: str, param_names: Iterable[str]) -> str:
    """Declare each ``param_names`` as a HIDDEN ReportParameter in an existing
    RDL if it isn't already -- so a parent's <Drillthrough> that forwards the
    value doesn't error "parameter not declared". Hidden + =Nothing default ->
    the standalone user never sees a prompt and "Refresh Fields" never asks.
    Inserts into <ReportParameters> (creating it just before <Body> in the
    RDL-2008 element order: ...EmbeddedImages, ReportParameters, Body...)."""
    todo = [p for p in (param_names or []) if p]
    if not todo:
        return rdl_text
    existing = {m.upper() for m in
                re.findall(r'<ReportParameter\s+Name="([^"]+)"', rdl_text)}
    todo = [p for p in todo if p.upper() not in existing]
    if not todo:
        return rdl_text
    blocks = "".join(
        f'<ReportParameter Name="{p}"><DataType>String</DataType>'
        f"<Nullable>true</Nullable><DefaultValue><Values><Value>=Nothing"
        f"</Value></Values></DefaultValue><AllowBlank>true</AllowBlank>"
        f"<Prompt>{p}</Prompt><Hidden>true</Hidden></ReportParameter>"
        for p in todo)
    if "<ReportParameters>" in rdl_text:
        return rdl_text.replace("</ReportParameters>",
                                blocks + "</ReportParameters>", 1)
    # No <ReportParameters> yet: create it immediately before <Body> (its
    # required position) -- match <Body ...> or <Body>.
    return re.sub(r"<Body(\s|>)",
                  "<ReportParameters>" + blocks + "</ReportParameters><Body" + r"\1",
                  rdl_text, count=1)


def _synth_report_from_sql(child_name: str, sql: str,
                           parent_param_names: Optional[Iterable[str]] = None,
                           drillthrough_params: Optional[Iterable[str]] = None):
    """Build a minimal ParsedReport from a SQL string so the real RDL
    generator and mockup renderer can produce a full child report. Columns
    come from the SELECT list; parameters from the SQL bind variables PLUS any
    ``drillthrough_params`` the parent forwards (so the child declares every
    parameter the parent's <Drillthrough> passes -- otherwise SSRS errors
    "the parameter 'X' is not declared" when the link is clicked).
    """
    from .models import ParsedReport, DataQuery, DataItem, ReportParameter
    safe = re.sub(r"[^A-Za-z0-9_]", "_", child_name or "SubReport") or "SubReport"
    rep = ParsedReport(name=child_name or "SubReport", dtd_version="(from artifacts)")
    cols = _select_columns(sql)
    items = [DataItem(name=c, label=c.replace("_", " ").title()) for c in cols]
    rep.queries.append(DataQuery(name=f"DS_{safe}", sql=sql or "",
                                 tsql=sql or "", items=items))
    pset = {p.upper() for p in (parent_param_names or [])}
    dt_set = {p.upper() for p in (drillthrough_params or [])}
    declared = set()
    for b in _bind_params_in_sql(sql):
        # A forwarded drill-through param stays HIDDEN even though the filter we
        # inject now binds it in the SQL -- the parent sets it; a standalone user
        # must never get a prompt box for it.
        visible = (b.upper() in pset) and (b.upper() not in dt_set)
        rep.parameters.append(
            ReportParameter(name=b, label=b, display=visible)
        )
        declared.add(b.upper())
    # Drill-through forwarded params (e.g. P_ORG_ID, P_SITE_ID). Declare each
    # even though the child SQL doesn't bind it: the parent's drill-through
    # forwards a value into it, and an undeclared target parameter is a hard
    # SSRS error. Hidden (display=False) -> the parent sets it, the standalone
    # user never sees an empty box; generate_rdl still gives it a =Nothing
    # default so "Refresh Fields" never prompts.
    for p in (drillthrough_params or []):
        if p and p.upper() not in declared:
            rep.parameters.append(ReportParameter(name=p, label=p, display=False))
            declared.add(p.upper())
    return rep


_INCLUSION_PARAM_NAME_RE = re.compile(r"(?i)site[_\s]?includ")


def _default_inclusion_params_yes(rdl_xml: str) -> str:
    """Default a child's 'site include' toggle parameter to 'YES'.

    A drill-through URL (``rs:Command=Render``) that omits a parameter renders
    the child with that parameter's DEFAULT. When a "site include" toggle is left
    at the neutral ``=Nothing`` default, an envelope-style child whose address +
    sort key both read ``DECODE(:P_Site_Include,'YES',S.Site_Name,NULL)`` falls to
    the NULL branch -- which BOTH drops the site from the address AND collapses the
    ORDER BY (keyed on that same column) to the recipient name. A "generate all"
    bulk envelope then prints in a DIFFERENT order than a parent run sorted by
    site, so the two stacks don't line up 1:1.

    Defaulting the toggle to ``YES`` makes the bulk/auto-rendered child sort by the
    site column, matching a site-sorted parent. Generic: keyed on the parameter
    NAME pattern (``site``+``includ``), never a report name; no-op when absent."""
    if not rdl_xml or "<ReportParameter" not in rdl_xml:
        return rdl_xml

    def _repl(m):
        block = m.group(0)
        if _INCLUSION_PARAM_NAME_RE.search(m.group(1)):
            return block.replace("<Value>=Nothing</Value>", "<Value>YES</Value>", 1)
        return block

    return re.sub(r'<ReportParameter Name="([^"]+)">.*?</ReportParameter>',
                  _repl, rdl_xml, flags=re.DOTALL)


_ENVELOPE_NAME_RE = re.compile(r"(?i)envelope|mailing|mail[_\s]?label|\blabel\b")
_ADDRESS_COL_RE = re.compile(r"(?i)addr")
_HARDCODED_TOP_RE = re.compile(
    r"\b([A-Za-z0-9][A-Za-z0-9.\-]{1,30})\s+is\s+hard\s?coded\s+at\s+the\s+top",
    re.IGNORECASE)


def _artifact_full_text(arts) -> str:
    """All readable text across the dropped artifacts (docx body + plain text).
    Used to mine layout notes the SQL alone doesn't carry -- e.g. a backend
    screenshot doc that says '<X> is hard coded at the top of the envelope'."""
    out = []
    for nm, blob in arts or []:
        try:
            low = (nm or "").lower()
            if low.endswith((".docx", ".dotx")):
                import io, zipfile
                z = zipfile.ZipFile(io.BytesIO(blob))
                for part in ("word/document.xml",):
                    if part in z.namelist():
                        out.append(re.sub(r"<[^>]+>", " ",
                                          z.read(part).decode("utf-8", "replace")))
            else:
                out.append(blob.decode("utf-8", "replace"))
        except Exception:
            continue
    return re.sub(r"\s+", " ", " ".join(out))


def _envelope_top_text(full_text: str) -> str:
    """Extract a hard-coded top-of-envelope label from an artifact note like
    'JV53063 is hard coded at the top of the envelope' -> 'JV 53063'. Generic:
    reads the value from the artifact, never a baked-in literal. '' if absent."""
    m = _HARDCODED_TOP_RE.search(full_text or "")
    if not m:
        return ""
    val = m.group(1).strip()
    # 'JV53063' -> 'JV 53063' (letters+digits run reads better spaced).
    sp = re.match(r"^([A-Za-z]+)(\d.*)$", val)
    return f"{sp.group(1)} {sp.group(2)}" if sp else val


def _address_column(cols) -> str:
    """The recipient-address column among the SELECT columns (name ~ 'addr'),
    else the longest-named candidate, else ''."""
    addr = [c for c in cols if _ADDRESS_COL_RE.search(c)]
    return addr[0] if addr else (cols[-1] if cols else "")


def _looks_like_envelope(child_name: str, cols) -> bool:
    """An envelope / mailing-label child: named like one AND carrying a
    recipient-address column. Conservative so ordinary list children are not
    re-routed to the positioned single-record envelope body."""
    named = bool(_ENVELOPE_NAME_RE.search(child_name or ""))
    has_addr = any(_ADDRESS_COL_RE.search(c) for c in (cols or []))
    return named and has_addr


def _attach_envelope_layout(rep, cols, top_text):
    """Give a (already param/query-wired) SQL child a POSITIONED mailing-envelope
    layout -- one envelope per recipient: the hard-coded top label centered near
    the top, the recipient address block in the window position. Replaces the
    labeled field-list with the envelope's real shape. Positions are
    envelope-archetype constants (where an address sits on an envelope); the LABEL
    + ADDRESS column come from the artifacts. Built ON TOP of the report
    _synth_report_from_sql already produced, so the drill-through param
    declarations + the :P_ORG_ID filter binding are preserved."""
    from .models import LayoutGroup, LayoutField
    qname = rep.queries[0].name if rep.queries else "Q_ADDRESS"
    addr = _address_column(cols)
    sec = LayoutGroup(name="section_main", kind="section_main",
                      repeat_on=qname, source_query=qname)
    frame = LayoutGroup(name="M_ENVELOPE", kind="frame",
                        x=0.0, y=0.0, width=7.9, height=6.0)
    fields = []
    if top_text:
        fields.append(LayoutField(name="B_TOP", kind="text", text=top_text,
                                  x=3.0, y=0.6, width=1.9, height=0.25,
                                  align="center", bold=True, font_size=11))
    if addr:
        fields.append(LayoutField(name="F_ADDRESS", kind="field", source=addr,
                                  x=3.6, y=4.2, width=4.0, height=1.3,
                                  align="start", font_size=11))
    frame.fields = fields
    sec.children.append(frame)
    rep.layout = [sec]
    return rep


def _strip_chrome_textboxes(rdl_xml: str) -> str:
    """Remove the fabricated 'Report run on:' + 'Page N of M' chrome textboxes
    the generator adds to a page header/footer -- a mailing envelope has neither.
    Keeps the page-title textbox (which carries the envelope's top label)."""
    for nm in ("Tb_RunOn", "Tb_PageNum"):
        rdl_xml = re.sub(r'<Textbox Name="' + nm + r'">.*?</Textbox>', "",
                         rdl_xml, flags=re.DOTALL)
    return rdl_xml


def _first_sql_from_paths(artifact_paths: Iterable[str]) -> str:
    """Pull the first runnable SELECT out of any artifact, trimmed to one
    statement (reuses the same extraction the stub builder uses)."""
    sql = ""
    for a in artifact_paths or []:
        if not os.path.isfile(a):
            continue
        candidate = _sql_from_artifact(a)
        if re.search(r"\bSELECT\b.+?\bFROM\b", candidate or "",
                     re.IGNORECASE | re.DOTALL):
            sql = candidate
            break
    if sql:
        sql = _trim_to_first_statement(sql)
    return sql


def _alias_outer_join_set(sql: str) -> set:
    """Aliases on the OUTER (nullable) side of an Oracle ``(+)`` join. Any column
    of such an alias is unsafe to equality-filter -- ``alias.col = :p`` silently
    discards the NULL-padded rows, turning the outer join into an inner one. We
    detect the alias from ANY ``alias.col(+)`` occurrence (the marker can sit on
    a different column of the same alias than the one we'd filter)."""
    return {m.group(1).upper()
            for m in re.finditer(r"\b([A-Za-z_]\w*)\.[A-Za-z_]\w*\s*\(\s*\+\s*\)",
                                 sql or "")}


def _inject_drillthrough_filter(sql: str,
                                drillthrough_params: Iterable[str]
                                ) -> Tuple[str, List[Tuple[str, str]]]:
    """Turn the child query's neutralized Oracle lexical filter (``&P_CRITERIA``)
    into a REAL, NULL-safe WHERE filter bound to the parent's forwarded
    drill-through params -- so a drilled link FILTERS to the clicked record while
    a standalone run (every param empty) still returns all rows.

    Generic, no report-specific names: for each forwarded ``P_<KEY>`` find a
    qualified id column whose base name == ``<KEY>`` (``P_ORG_ID`` -> ``O.Org_Id``),
    preferring an INNER-joined column (its alias has no ``(+)`` marker) so the
    filter never breaks an outer join. The fragment is spliced in at the lexical
    slot (where the report author intended the dynamic WHERE); any leftover
    lexical is commented out. Returns ``(new_sql, applied)`` where ``applied`` is
    the ``[(param, column), ...]`` actually wired (empty -> nothing matched, the
    caller keeps the unfiltered behaviour + the manual-wiring note)."""
    dt = [p for p in (drillthrough_params or []) if p]
    if not sql or not dt:
        return sql, []
    lexicals = _lexical_refs_in_sql(sql)
    if not lexicals:
        return sql, []  # no designated filter slot -> don't guess a WHERE
    id_cols = _id_columns_in_sql(sql)
    outer = _alias_outer_join_set(sql)

    def _col_for(param: str):
        key = re.sub(r"(?i)^p_", "", param).upper()  # P_ORG_ID -> ORG_ID
        for c in id_cols:  # appearance order == FROM order
            alias, base = c.split(".")[0].upper(), c.split(".")[-1].upper()
            if base == key and alias not in outer:
                return c  # first inner-joined match wins
        return None  # only outer-joined matches -> skip (would break the join)

    applied: List[Tuple[str, str]] = []
    used_cols = set()
    for p in dt:
        c = _col_for(p)
        if c and c.upper() not in used_cols:
            applied.append((p, c))
            used_cols.add(c.upper())
    if not applied:
        return sql, []
    frag = "\n\t" + "\n\t".join(
        f"AND (:{p} IS NULL OR {c} = :{p})" for p, c in applied) + "\n"
    first = lexicals[0]
    new_sql = re.sub(r"&" + re.escape(first) + r"\b", frag, sql, count=1)
    # Any remaining lexical refs -> runnable comment (no SSRS equivalent).
    new_sql = re.sub(r"&([A-Za-z_][A-Za-z0-9_]*)",
                     r"/* lexical ref &\1 -- neutralized (no SSRS equivalent) */",
                     new_sql)
    return new_sql, applied


def _drillthrough_reconciliation_notes(sql: str,
                                       drillthrough_params: Iterable[str],
                                       applied_filter: Optional[
                                           List[Tuple[str, str]]] = None
                                       ) -> List[str]:
    """Human-readable guidance for the drill-through child's filtering. The
    parent forwards ``drillthrough_params`` (e.g. P_ORG_ID, P_SITE_ID); we
    always declare them so the link can't error. If ``applied_filter`` is set
    the child query was rewritten to FILTER by those params (success); otherwise
    the child runs UNFILTERED and we tell the user exactly how to wire it.
    Generic -- candidate columns are read from the child SQL."""
    notes: List[str] = []
    dt = [p for p in (drillthrough_params or []) if p]
    if not dt:
        return notes
    notes.append(
        "Declared drill-through target parameter(s) " + ", ".join(dt) +
        " (forwarded by the parent's link) so the drill-through won't error "
        "with \"parameter not declared\"."
    )
    if applied_filter:
        notes.append(
            "The drill-through link now FILTERS this child to the clicked "
            "record: " + "; ".join(f"{p} = {c}" for p, c in applied_filter) +
            " (added a NULL-safe \"AND (:param IS NULL OR col = :param)\" to the "
            "query at the lexical filter slot). Clicking a parent row opens ONLY "
            "that record's page here; running this child on its own (parameters "
            "left empty) still lists every record."
        )
        return notes
    lex = _lexical_refs_in_sql(sql)
    id_cols = _id_columns_in_sql(sql)
    if lex:
        cand = (" Candidate key columns in this query: " +
                ", ".join(id_cols[:8]) + ".") if id_cols else ""
        notes.append(
            "This child filters through the Oracle lexical " +
            ", ".join("&" + x for x in lex) + " (a runtime SQL fragment with no "
            "SSRS equivalent; it was neutralized so the query is valid and "
            "runs UNFILTERED). To make the link filter to the drilled record, "
            "open the dataset query and add a WHERE condition that binds the "
            "forwarded param(s), e.g.  AND <key_column> = :" + dt[0] +
            (("  AND <key_column> = :" + dt[1]) if len(dt) > 1 else "") + "." +
            cand
        )
    return notes


def build_subreport(child_name: str,
                    artifact_paths: Iterable[str],
                    parent_param_names: Optional[Iterable[str]] = None,
                    drillthrough_params: Optional[Iterable[str]] = None
                    ) -> Dict[str, Any]:
    """Build a child report from artifacts and return a rich preview payload.

    Returns a dict with:
        rdl_xml, mockup_html, mockup_backend_html,
        fields, binds, forwarded_params, sql, issues, source, report_name

    ``drillthrough_params`` are the parameter names the PARENT forwards to this
    child via its drill-through link (see ``forwarded_drillthrough_params``).
    The child RDL declares each one so the parent's <Drillthrough> resolves --
    an undeclared target parameter is a hard SSRS error the instant the link is
    clicked.

    ``source`` is one of "oracle_xml", "rdl", "sql", "stub" -- which branch
    produced the output, so the UI can label the preview honestly.
    """
    issues: List[str] = []
    artifact_paths = list(artifact_paths or [])
    drillthrough_params = list(drillthrough_params or [])
    arts = _read_artifacts(artifact_paths)
    safe_name = re.sub(r"[^A-Za-z0-9_]", "_", child_name or "SubReport") or "SubReport"

    def _render(rep, mode="frontend") -> str:
        try:
            from .preview.html_mockup import render_mockup
            return render_mockup(rep, mode=mode)
        except Exception as e:  # noqa: BLE001
            return f"<em>Preview unavailable: {type(e).__name__}: {e}</em>"

    # 1. Oracle Reports XML present -> full pipeline (best fidelity).
    for nm, blob in arts:
        if _looks_like_oracle_xml(nm, blob):
            try:
                from . import convert as _convert
                # Forward the parent's drill-through params so the child RDL
                # DECLARES them (else the link errors when clicked). Previously
                # the Oracle-XML path -- the highest-fidelity, most common one
                # -- silently dropped them.
                data = _convert(blob, extra_param_names=drillthrough_params)
                rdl_xml = data.get("rdl_xml", "")
                # Columns we surfaced = the RDL dataset field names.
                fields = re.findall(r'<Field Name="([^"]+)"', rdl_xml)
                params = [p["name"] for p in
                          (data.get("report") or {}).get("parameters", [])]
                dt_upper = {p.upper() for p in drillthrough_params}
                forwarded = [p for p in params if p.upper() in dt_upper]
                return {
                    "rdl_xml": rdl_xml,
                    "mockup_html": data.get("mockup_html", ""),
                    "mockup_backend_html": data.get("mockup_backend_html", ""),
                    "fields": fields,
                    "binds": params,
                    "forwarded_params": forwarded,
                    "sql": "",
                    "issues": issues,
                    "source": "oracle_xml",
                    "report_name": (data.get("report") or {}).get("name") or child_name,
                }
            except Exception as e:  # noqa: BLE001
                issues.append(f"Oracle XML found but full conversion failed "
                              f"({type(e).__name__}: {e}); trying other artifacts.")

    # 2. Existing SSRS .rdl present -> use it as-is; preview from its dataset.
    for nm, blob in arts:
        if _looks_like_rdl(nm, blob):
            rdl_text = blob.decode("utf-8", "replace")
            if "<Report" in rdl_text:
                # Declare the parent's forwarded drill-through params in the
                # supplied RDL if it doesn't already -- else the link errors
                # when clicked (same gap the Oracle-XML path had).
                rdl_text = _inject_report_parameters(rdl_text, drillthrough_params)
                rep = _report_from_rdl(rdl_text, child_name)
                main = rep.queries[0] if rep.queries else None
                forwarded = [p for p in drillthrough_params
                             if f'Name="{p}"' in rdl_text]
                return {
                    "rdl_xml": rdl_text,
                    "mockup_html": _render(rep, "frontend"),
                    "mockup_backend_html": _render(rep, "backend"),
                    "fields": [i.name for i in (main.items if main else [])],
                    "binds": [],
                    "forwarded_params": forwarded,
                    "sql": (main.sql if main else ""),
                    "issues": issues + [
                        "Used the supplied .rdl as-is; preview derived from its "
                        "largest dataset."
                    ],
                    "source": "rdl",
                    "report_name": rep.name,
                }

    # 3. SQL-bearing artifacts (.sql/.docx/.txt) -> synth report -> real RDL.
    sql = _first_sql_from_paths(artifact_paths)
    applied_filter: List[Tuple[str, str]] = []
    if sql:
        # Wire the parent's forwarded params into the child's WHERE so the link
        # FILTERS to the drilled record (the parameterized envelope the user
        # wants), instead of always opening the unfiltered first page.
        sql, applied_filter = _inject_drillthrough_filter(sql, drillthrough_params)
        _cols = _select_columns(sql)
        rep = _synth_report_from_sql(child_name, sql, parent_param_names,
                                     drillthrough_params)
        _envelope = _looks_like_envelope(child_name, _cols)
        if _envelope:
            # Mailing-envelope child: give the (param/query-wired) report a
            # POSITIONED envelope layout (top label + address in the window)
            # instead of the labeled field-list. The top label is mined from the
            # artifacts (e.g. a screenshot note).
            top_text = _envelope_top_text(_artifact_full_text(arts))
            _attach_envelope_layout(rep, _cols, top_text)
            if top_text:
                issues.append(f"Envelope archetype: positioned the address block "
                              f"+ top label '{top_text}' (from your artifacts).")
            else:
                issues.append("Envelope archetype: positioned the address block "
                              "(no hard-coded top label found in the artifacts).")
        try:
            from .translators.plsql_to_tsql import translate_report
            translate_report(rep)
        except Exception as e:  # noqa: BLE001
            issues.append(f"Translator warning: {type(e).__name__}: {e}")
        rdl_xml = ""
        try:
            from .generators.rdl import generate_rdl
            rdl_xml = generate_rdl(rep)
        except Exception as e:  # noqa: BLE001
            issues.append(f"RDL generation failed ({type(e).__name__}: {e}); "
                          f"used minimal stub.")
            rdl_xml = compose_subreport_rdl(
                safe_name, artifact_paths, parent_param_names,
                drillthrough_params)["rdl_xml"]
        if _envelope:
            # An envelope has no 'Report run on' / 'Page of' chrome.
            rdl_xml = _strip_chrome_textboxes(rdl_xml)
        # A bulk "generate all" link auto-renders the child with its parameter
        # DEFAULTS; default a site-include toggle to YES so the bulk envelopes
        # sort by site, matching a site-sorted parent (the 1:1 mailing order).
        rdl_xml = _default_inclusion_params_yes(rdl_xml)
        main = rep.queries[0] if rep.queries else None
        cols = [i.name for i in (main.items if main else [])]
        binds = _bind_params_in_sql(sql)
        pset = {p.upper() for p in (parent_param_names or [])}
        # "Forwarded" = every parameter the child now declares that the parent
        # passes: SQL binds the parent declares + the drill-through URL params.
        dt_upper = {p.upper() for p in drillthrough_params}
        forwarded = [b for b in binds if b.upper() in pset]
        for p in drillthrough_params:
            if p and p.upper() not in {f.upper() for f in forwarded}:
                forwarded.append(p)
        if not cols:
            issues.append("Could not infer a column list from the SQL; the "
                          "preview shows generic columns -- refine in Report Builder.")
        issues.extend(_drillthrough_reconciliation_notes(
            sql, drillthrough_params, applied_filter))
        if pset and binds and not forwarded:
            issues.append("Child SQL has bind variables but none match the "
                          "parent's parameters -- drill-through will not forward values.")
        return {
            "rdl_xml": rdl_xml,
            "mockup_html": _render(rep, "frontend"),
            "mockup_backend_html": _render(rep, "backend"),
            "fields": cols,
            "binds": binds,
            "forwarded_params": forwarded,
            "sql": sql,
            "issues": issues,
            "source": "sql",
            "report_name": rep.name,
        }

    # 4. Nothing parseable -> stub + a friendly "drop more" note.
    stub = compose_subreport_rdl(safe_name, artifact_paths, parent_param_names,
                                 drillthrough_params)
    note = (
        "<div style=\"font-family:Segoe UI,Arial,sans-serif;padding:20px;"
        "color:#333;line-height:1.5\">"
        "<div style=\"font-weight:700;font-size:15px;margin-bottom:6px\">"
        "Nothing to preview yet</div>"
        "Drop the child report's <b>Oracle XML</b>, an existing <b>.rdl</b>, or "
        "its <b>SQL</b> (.sql / .docx / .txt) and the preview will render here, "
        "just like the main report.</div>"
    )
    return {
        "rdl_xml": stub["rdl_xml"],
        "mockup_html": note,
        "mockup_backend_html": note,
        "fields": stub.get("fields", []),
        "binds": stub.get("binds", []),
        "forwarded_params": stub.get("forwarded_params", []),
        "sql": stub.get("sql", ""),
        "issues": issues + stub.get("issues", []),
        "source": "stub",
        "report_name": child_name or safe_name,
    }


__all__ = [
    "detect_subreport_links",
    "is_drillthrough_only",
    "compose_subreport_rdl",
    "build_subreport",
    "forwarded_drillthrough_params",
]
